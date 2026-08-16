import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

files_description = {
    "lib/firebase_options.dart": "ملف تكوين إعدادات Firebase الخاصة بالمشروع (الإعدادات التلقائية لكل منصة مثل Android و iOS).",
    "lib/main.dart": "نقطة الدخول الرئيسية للتطبيق، حيث يتم تهيئة الإعدادات الأولية وإطلاق التطبيق.",
    "lib/models/event.dart": "نموذج البيانات (Model) الخاص بالفعاليات (Events). يحتوي على خصائص الفعالية مثل الاسم، التاريخ، الموقع، إلخ.",
    "lib/models/models.dart": "ملف لتصدير جميع النماذج (Models) لتسهيل استيرادها في باقي أجزاء التطبيق.",
    "lib/models/reservation.dart": "نموذج البيانات الخاص بحجوزات المستخدمين.",
    "lib/models/review.dart": "نموذج البيانات الخاص بالتقييمات والمراجعات.",
    "lib/models/user.dart": "نموذج البيانات الخاص بالمستخدمين.",
    "lib/providers/admin_provider.dart": "مزود الحالة (Provider) الخاص بإدارة وظائف لوحة تحكم الإدارة (Admin).",
    "lib/providers/auth_provider.dart": "مزود الحالة الخاص بالمصادقة (تسجيل الدخول، إنشاء الحساب، وتسجيل الخروج).",
    "lib/providers/event_provider.dart": "مزود الحالة الخاص ببيانات الفعاليات (لجلب الفعاليات وإدارتها).",
    "lib/providers/notification_provider.dart": "مزود الحالة لإدارة وتتبع الإشعارات الواردة للمستخدم.",
    "lib/providers/providers.dart": "ملف تجميع وتصدير لجميع مزودي الحالة (Providers).",
    "lib/providers/reservation_provider.dart": "مزود الحالة الخاص بخدمة الحجوزات وإدارتها.",
    "lib/providers/theme_provider.dart": "مزود الحالة الخاص بإدارة سمة التطبيق (الوضع الليلي والنهاري).",
    "lib/services/admin_service.dart": "خدمة (Service) تتواصل مع قاعدة البيانات لتنفيذ مهام مدير النظام (Admin).",
    "lib/services/auth_service.dart": "خدمة معالجة أوامر المصادقة والتواصل مع Firebase Authentication.",
    "lib/services/event_service.dart": "خدمة للتعامل مع العمليات الخاصة بالفعاليات (إنشاء، جلب، تعديل، حذف).",
    "lib/services/payment_service.dart": "خدمة لمعالجة الدفع وعمليات الشراء داخل التطبيق.",
    "lib/services/reservation_service.dart": "خدمة لمعالجة عمليات الحجز وتحديثات حالة الحجوزات.",
    "lib/services/review_service.dart": "خدمة للتعامل مع إضافة وجلب المراجعات والتقييمات الخاصة بالفعاليات.",
    "lib/services/services.dart": "ملف تجميع وتصدير لجميع الخدمات (Services).",
    "lib/utils/app_theme.dart": "يحتوي على تعريفات سمات التطبيق (الألوان، الخطوط، أنماط العرض المخصصة).",
    "lib/views/admin/admin_dashboard.dart": "واجهة لوحة التحكم الخاصة بالمدير لعرض الإحصائيات والإدارة العامة.",
    "lib/views/admin/admin_notification_page.dart": "واجهة مخصصة لمدير النظام لإدارة وإرسال الإشعارات.",
    "lib/views/admin/admin_profile_page.dart": "واجهة الملف الشخصي الخاص بمدير النظام.",
    "lib/views/admin/manage_user_dialog.dart": "نافذة منبثقة (Dialog) لمدير النظام لإدارة حالة أو معلومات مستخدم معين.",
    "lib/views/auth/login_page.dart": "واجهة تسجيل الدخول للمستخدمين للحصول على صلاحيات الدخول.",
    "lib/views/calendar/calendar_page.dart": "واجهة لتمثيل الفعاليات على شكل تقويم زمني لتسهيل التصفح.",
    "lib/views/home/event_list_page.dart": "الواجهة الرئيسية التي تعرض قائمة الفعاليات المتاحة للمستخدم.",
    "lib/views/map/event_map_page.dart": "واجهة تعرض خريطة تفاعلية توضح مواقع الفعاليات جغرافياً.",
    "lib/views/notifications/notification_page.dart": "واجهة لعرض قائمة الإشعارات التي يتلقاها المستخدم.",
    "lib/views/organizer/create_event_page.dart": "واجهة مخصصة للمنظمين (Organizers) لإنشاء تفاصيل فعالية جديدة.",
    "lib/views/organizer/event_reservations_page.dart": "واجهة للمنظمين لعرض ومتابعة الحجوزات الخاصة بفعاليتهم.",
    "lib/views/organizer/organizer_dashboard.dart": "لوحة تحكم رئيسية مصغرة خاصة بمنظم الفعاليات.",
    "lib/views/organizer/organizer_stats_page.dart": "واجهة تستعرض إحصائيات حول أداء الفعاليات للمنظم.",
    "lib/views/reservation/reservation_confirmation_page.dart": "واجهة تأكيد عملية الحجز ونجاحها.",
    "lib/views/user/edit_profile_dialog.dart": "نافذة منبثقة للمستخدم لتعديل معلومات حسابه الشخصي.",
    "lib/views/user/event_detail_page.dart": "واجهة عرض تفاصيل فعالية معينة بنظرة شاملة.",
    "lib/views/user/my_reservations_page.dart": "واجهة مستخدم توضح حجوزاته الحالية والسابقة.",
    "lib/views/user/profile_page.dart": "واجهة الملف الشخصي للمستخدم لعرض الإعدادات وبيانات الحساب.",
    "lib/widgets/event_card.dart": "مكون واجهة مستخدم (Widget) لعرض نبذة مصغرة وبطاقة تفصيلية عن الفعالية.",
    "lib/widgets/glass_card.dart": "مكون لإنشاء بطاقة بتأثير الزجاج (Glassmorphism) لعمل تصميم جمالي (Premium UI).",
    "lib/widgets/location_picker_dialog.dart": "مكون منبثق لاختيار الموقع عبر الخريطة.",
    "lib/widgets/notification_badge.dart": "مكون لعرض دائرة صغيرة تحتوي على عدد الإشعارات الجديدة فوق الأيقونة.",
    "lib/widgets/payment_dialog.dart": "نافذة منبثقة لإدخال أو معالجة معلومات الدفع.",
    "lib/widgets/premium_button.dart": "زر بتصميم متميز (Premium) لاستخدامه في أجزاء مهمة من الواجهة.",
    "lib/widgets/premium_hero_background.dart": "مكون واجهة لعرض خلفية متميزة وجذابة (Hero) في أعلى الشاشات.",
    "lib/widgets/review_dialog.dart": "نافذة منبثقة لإعطاء تقييم للفعالية كتابةً.",
    "lib/widgets/shimmer_card.dart": "مكون واجهة يعرض تأثير التحميل (Shimmer Effect) قبل ظهور البيانات.",
    "lib/widgets/widgets.dart": "ملف تجميع وتصدير لجميع المكونات (Widgets) المخصصة في التطبيق لسهولة مشاركتها.",
    "lib/widgets/admin/admin_event_tile.dart": "مكون مخصص للمشرفين لعرض الفعالية بشكل مبسط في قائمة.",
    "lib/widgets/admin/admin_search_bar.dart": "شريط بحث مخصص لاستخدام المشرفين لإيجاد مستخدمين أو فعاليات.",
    "lib/widgets/admin/admin_user_tile.dart": "مكون لعرض مستخدم واحد بشكل مبسط ضمن لوحة تحكم الإدارة.",
    "lib/widgets/admin/admin_widgets.dart": "تصدير لمكونات الواجهة الخاصة بقسم الإدارة.",
    "lib/widgets/admin/kpi_card.dart": "بطاقة لعرض مؤشرات الأداء الرئيسية (KPIs) لمدير النظام (كـ إجمالي المبيعات، الحجوزات، والأعداد)."
}

document = Document()

# Set right-to-left alignment for the whole document
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# Set RTL property for the style
style.font.rtl = True

# Heading
heading = document.add_heading('شرح ملفات نظام التطبيق والمكونات الخاصة به', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Intro paragraph
intro = document.add_paragraph('فيما يلي شرح تفصيلي لدور كل ملف في مستودع الكود المصدري (داخل مجلد lib):')
intro.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Table
table = document.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[1].text = 'مسار الملف'
hdr_cells[0].text = 'الدور / الشرح'

for cell in hdr_cells:
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            run.font.bold = True

for path, desc in files_description.items():
    row_cells = table.add_row().cells
    
    # Text
    p0 = row_cells[0].paragraphs[0]
    p0.text = desc
    p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Path
    p1 = row_cells[1].paragraphs[0]
    p1.text = path
    # Even if path is ltr, in RTL tables it might look better aligned left or right, let's just make it right.
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
document.save('files_explanations_arabic.docx')
print("Document saved successfully.")
